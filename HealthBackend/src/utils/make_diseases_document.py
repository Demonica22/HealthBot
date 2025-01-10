import docx
from io import BytesIO
from src.utils.localization import constraints


def make_in_memory_document(diseases: list[dict], user_language: str):
    doc = docx.Document()
    doc.add_heading(constraints[user_language]['table_title'], level=1)
    headers = [constraints[user_language]['title'],
               constraints[user_language]['description'],
               constraints[user_language]['treatment_plan'],
               constraints[user_language]['date_from'],
               constraints[user_language]['date_to'],
               constraints[user_language]['still_sick'], ]
    table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')
    table.autofit = True
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
    for illness in diseases:
        row_cells = table.add_row().cells
        row_cells[0].text = illness["title"]
        row_cells[1].text = illness["description"]
        row_cells[2].text = illness["treatment_plan"]
        row_cells[3].text = illness["date_from"].strftime("%Y-%m-%d")
        row_cells[4].text = illness["date_to"].strftime("%Y-%m-%d") if illness["date_to"] else "-"
        row_cells[5].text = constraints[user_language]['yes'] if illness["still_sick"] else constraints[user_language][
            'no']

    # Save the document to memory
    memory_stream = BytesIO()
    doc.save(memory_stream)
    memory_stream.seek(0)

    return memory_stream
